"""
Office and document operations.

Handles Excel, Word, PDF, and other document formats.
"""

import asyncio
import logging
import os
from pathlib import Path
from typing import Optional

import config

log = logging.getLogger(__name__)


async def excel_read(path: str, sheet: str = None) -> str:
    """Read Excel file and return contents as text."""
    try:
        import openpyxl
        import pandas as pd

        path_obj = Path(path).resolve()

        if not path_obj.exists():
            return f"[error] File not found: {path}"

        # Try with pandas for simplicity
        if sheet:
            df = pd.read_excel(path_obj, sheet_name=sheet)
        else:
            df = pd.read_excel(path_obj)

        # Convert to text
        output = f"## {path_obj.name}\n"
        output += f"Shape: {df.shape[0]} rows x {df.shape[1]} columns\n\n"
        output += df.to_string(max_rows=50, max_colwidth=50)

        return output[:8000]

    except Exception as e:
        return f"[error] {e}"


async def excel_write(path: str, data: str, sheet: str = "Sheet1") -> str:
    """Write data to Excel file."""
    try:
        import pandas as pd

        path_obj = Path(path).resolve()
        path_obj.parent.mkdir(parents=True, exist_ok=True)

        # Try to parse as CSV-like data
        lines = data.strip().split("\n")
        if len(lines) < 2:
            return "[error] Need at least header + 1 row"

        # Assume first line is headers
        import io

        df = pd.read_csv(io.StringIO("\n".join(lines)))

        if sheet:
            with pd.ExcelWriter(path_obj, engine="openpyxl") as writer:
                df.to_excel(writer, sheet_name=sheet, index=False)
        else:
            df.to_excel(path_obj, index=False)

        return f"[excel] Written {len(df)} rows to {path}"

    except Exception as e:
        return f"[error] {e}"


async def excel_append(path: str, data: str, sheet: str = "Sheet1") -> str:
    """Append rows to existing Excel file."""
    try:
        import openpyxl
        import pandas as pd

        path_obj = Path(path).resolve()

        # Read existing
        if path_obj.exists():
            existing = pd.read_excel(path_obj, sheet_name=sheet)
        else:
            existing = pd.DataFrame()

        # Parse new data
        import io

        new_data = pd.read_csv(io.StringIO(data))

        # Append
        combined = pd.concat([existing, new_data], ignore_index=True)

        with pd.ExcelWriter(path_obj, engine="openpyxl") as writer:
            combined.to_excel(writer, sheet_name=sheet, index=False)

        return f"[excel] Appended {len(new_data)} rows, total {len(combined)}"

    except Exception as e:
        return f"[error] {e}"


async def excel_formulas(path: str, formula: str, cell: str = "A1") -> str:
    """Add formula to a cell."""
    try:
        import openpyxl

        path_obj = Path(path).resolve()

        if not path_obj.exists():
            return "[error] File not found"

        wb = openpyxl.load_workbook(path_obj)
        ws = wb.active

        ws[cell] = formula
        wb.save(path_obj)

        return f"[excel] Added formula to {cell}: {formula}"

    except Exception as e:
        return f"[error] {e}"


async def word_read(path: str) -> str:
    """Read Word document."""
    try:
        from docx import Document

        path_obj = Path(path).resolve()

        if not path_obj.exists():
            return f"[error] File not found: {path}"

        doc = Document(path_obj)

        output = []
        for para in doc.paragraphs:
            if para.text.strip():
                output.append(para.text)

        return "\n".join(output)[:8000]

    except Exception as e:
        return f"[error] {e}"


async def word_write(path: str, content: str) -> str:
    """Write to Word document."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt

        path_obj = Path(path).resolve()
        path_obj.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()

        # Add content
        for line in content.split("\n"):
            if line.strip():
                doc.add_paragraph(line)

        doc.save(path_obj)

        return f"[word] Written to {path}"

    except Exception as e:
        return f"[error] {e}"


async def pdf_read(path: str, page: int = None) -> str:
    """Read PDF text content."""
    try:
        from pypdf import PdfReader

        path_obj = Path(path).resolve()

        if not path_obj.exists():
            return f"[error] File not found: {path}"

        reader = PdfReader(path_obj)

        if page is not None:
            if page < 0 or page >= len(reader.pages):
                return f"[error] Page {page} out of range (0-{len(reader.pages) - 1})"
            text = reader.pages[page].extract_text()
            return f"## Page {page + 1}\n{text}"[:8000]

        # All pages
        output = []
        for i, p in enumerate(reader.pages):
            text = p.extract_text()
            if text.strip():
                output.append(f"--- Page {i + 1} ---\n{text}")

        return "\n\n".join(output)[:8000]

    except Exception as e:
        return f"[error] {e}"


async def pdf_info(path: str) -> str:
    """Get PDF metadata."""
    try:
        from pypdf import PdfReader

        path_obj = Path(path).resolve()

        if not path_obj.exists():
            return f"[error] File not found: {path}"

        reader = PdfReader(path_obj)
        meta = reader.metadata

        output = [f"## {path_obj.name}"]
        output.append(f"Pages: {len(reader.pages)}")

        if meta:
            if meta.get("/Title"):
                output.append(f"Title: {meta['/Title']}")
            if meta.get("/Author"):
                output.append(f"Author: {meta['/Author']}")
            if meta.get("/Creator"):
                output.append(f"Creator: {meta['/Creator']}")

        return "\n".join(output)

    except Exception as e:
        return f"[error] {e}"


async def image_info(path: str) -> str:
    """Get image dimensions and info."""
    try:
        from PIL import Image

        path_obj = Path(path).resolve()

        if not path_obj.exists():
            return f"[error] File not found: {path}"

        img = Image.open(path_obj)

        info = [
            f"## {path_obj.name}",
            f"Format: {img.format}",
            f"Size: {img.size[0]} x {img.size[1]} pixels",
            f"Mode: {img.mode}",
        ]

        if hasattr(img, "info"):
            if img.info.get("dpi"):
                info.append(f"DPI: {img.info['dpi']}")

        return "\n".join(info)

    except Exception as e:
        return f"[error] {e}"
